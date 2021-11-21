from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def add_clan(clan, document, *data):
    """
    This method should be used for creating paragraph for given data and document
    :param clan:
    :param document:
    :param data:
    :return:
    """
    t = document.add_heading(template_1_c[clan], level=3)
    t.bold = True
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = t.add_run()
    run.add_break()
    p = document.add_paragraph(template_1_p[clan].format(*data)) if data \
        else document.add_paragraph(template_1_p[clan])
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = p.add_run()
    run.add_break()


# titles of paragraphs
template_1_c = {
    "clan_1": "Član 1.",
    "clan_2": "Član 2.",
    "clan_3": "Član 3.",
    "clan_4": "Član 4.",
    "clan_5": "Član 5.",
    "clan_6": "Član 6.",
}

# texts of paragraphs
template_1_p = {
    "clan_1": "Prodavac je jedini vlasnik nepokretnosti koja se sastoji od useljivog stana koji se nalazi na adresi "
        "{}, na {} spratu, stan broj {}, površine "
        "{} m2, po strukturi __ kat.parcela ______ KO"
        " _____ upisan/a u listu nepokretnosti br. __. Prodavac prodaje Kupcu nepokretnost opisanu u stavu 1. "
        "ovog Ugovora u viđenom stanju za kupoprodajnu cenu koja iznosi {} EUR. "
        "Prodavac je stekao opisanu nepokretnost na osnovu Ugovora o otkupu Ov. br. __ od __ godine "
        "kod __ suda u __ na osnovu kupoprodajnog ugovora Ov.br. __ od __ kod __ suda u __ .",
    "clan_2": f'Prodavac potvrđuje da je primio od Kupca u celosti ugovorenu cenu iz člana 1. ovog Ugovora. '
        f'Alternativa : Kupac isplaćuje Prodavcu prilikom zaključivanja i overe ovog Ugovora iznos od '
        f'________ EUR kao ostatak kupoprodajne cene obzirom da je Kupac na osnovu zaključenog '
        f'Predugovora o kupoprodaji dana __ godine isplatio Prodavcu iznos od __ EUR kao kaparu. '
        f'Porez na promet apsolutnih prava kao i sve takse u ovom pravnom poslu snosi Kupac.',
    "clan_3": f'Kupac potvrđuje da je od Prodavca primio u posed i državinu nepokretnost iz člana 1. ovog Ugovora '
        f'na dan potpisivanja ovog Ugovora kao i da je primio od Prodavca sve dokaze o njegovom pravu vlasništva '
        f'na nepokretnosti. Alternativa : Prodavac se obavezuje da nepokretnost iz člana 1. ovog Ugovora isprazni'
        f' od lica i stvari najkasnije do ___. godine i da je, u stanju u kome se ona nalazi u vreme zaključenja ovog'
        f' Ugovora do __ godine preda Kupcu u posed i državinu kao i da Kupcu preda sve dokaze svog prava vlasništva '
        f'na nepokretnosti.',
    "clan_4": f'Prodavac je potpisom ovog Ugovora saglasan da Kupac može u katastru nepokretnosti uknjižiti svoje pravo '
        f'vlasništva na nepokretnosti bliže opisanoj u članu 1. ovog Ugovora bez ikakve njegove dalje pismene ili '
        f'usmene saglasnosti tj. daje Kupcu neopozivu “clausulu intabulandi”.',
    "clan_5": f'Prodavac pod punom krivičnom i materijalnom odgovornošću izjavljuje i garantuje Kupcu da je isključivi '
        f'vlasnik opisane nepokretnosti, da nepokretnost nije opterećena nikakvim upisanim ili prećutnim teretima '
        f'(zaloge, hipoteke i dr.) niti pravima trećih lica, da ista nije predmet sudskog ili upravnog postupka, '
        f'da nije predmet podele bračne tekovine ili nasleđivanja, da nema zabrane raspolaganja ili otuđenja, '
        f'da nije predmet druge kupoprodaje, da ne postoji smetnja da se izvrši prenos vlasništva i predaja '
        f'nepokretnosti u posed, da nije za nju već primio kaparu ili naknadu, da nije data u zakup ili na poklon, '
        f'da nije predmet ugovora o doživotnom izdržavanju, ugovora o zameni, da nije predmet bilo kog drugog pravnog '
        f'posla i obavezuje se da kupcu pruži kompletnu zaštitu od pravnog uznemiravanja – „evikcije“ pod pretnjom'
        f' naknade štete i raskida ovog Ugovora. Ukoliko Prodavac ne bude uspeo da zaštiti Kupca od eventualnih '
        f'prava trećih lica na nepokretnosti koji je predmet ove kupoprodaje i Kupac bude primoran da se iseli iz '
        f'nepokretnosti, Prodavac, potpisom ovog Ugovora neopozivo izjavljuje da će Kupcu u roku od 15 dana od dana '
        f'njegovog iseljenja, izvršiti povraćaj celokupnog iznosa kupoprodajne cene u valuti u kojoj je isplaćena '
        f'kupoprodajna cena kao i da će Kupcu naknaditi svu štetu.',
    "clan_6": f'Sve troškove oko potpisivanja ovog Ugovora i njegove overe snosi Kupac. Prodavac je obavezan da izmiri '
        f'sve obaveze za komunalije za opisanu nepokretnost (porez na imovinu, električnu energiju, vodu,'
        f' telefon i ostalo) zaključno sa mesecom isplate ostatka celokupne kupoprodajne cene.',
    "clan_7": f'U slučaju spora nadležan je sud u __. Ugovorne strane su poučene o važnosti ovog Ugovora i potvrđuju da im '
       f'je ovaj Ugovor pročitan, da ga slobodnom voljom prihvataju i saglasno tome potpisuju. Ugovor je sačinjen '
       f'u __ (__) istovetnih primeraka od kojih svaka ugovorna strana zadržava po __ (__) primerka a '
       f'završen je zaključno sa ovom klauzulom.'
}
