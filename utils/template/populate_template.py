import io

from docx import Document
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from utils.template import template_helper


def create_template_ugovor(data):
    """
    This method should be used for creating template "ugovor" for "pkupac"
    :param data:
    :return:
    """
    # create document
    document = Document()

    # add title
    title = document.add_heading('UGOVOR O KUPOPRODAJI STANA', level=2)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.add_run()
    run.add_break()

    # add description for document
    document.add_paragraph(f'Koji zaključuju dana {datetime.now().date()} godine ugovorne strane :')
    document.add_paragraph(f'Prodavac: Adresa ##### JMBG #####')
    document.add_paragraph(
        f'Kupac: Adresa {data.kupac.adresa} JMBG {data.kupac.jmbg}'
    )

    # add paragraphs
    template_helper.add_clan(
        'clan_1',
        document,
        data.stan.adresa,
        data.stan.sprat,
        data.stan.broj_stana,
        data.stan.kvadratura,
        data.cena_za_kupca
    )
    template_helper.add_clan(clan='clan_2', document=document)
    template_helper.add_clan(clan='clan_3', document=document)
    template_helper.add_clan(clan='clan_4', document=document)
    template_helper.add_clan(clan='clan_5', document=document)
    template_helper.add_clan(clan='clan_6', document=document)

    f1 = document.add_paragraph(f'Kupac: {data.kupac.first_name} {data.kupac.last_name}')
    f1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    f2 = document.add_paragraph(f'________________________')
    f2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    document.save(f'{data.kupac.first_name}_{data.kupac.last_name}.docx')

    # Create in-memory buffer
    file_stream = io.BytesIO()
    # Save the .docx to the buffer
    document.save(file_stream)
    # Reset the buffer's file-pointer to the beginning of the file
    file_stream.seek(0)

    return file_stream
